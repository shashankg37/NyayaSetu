from datetime import date
from pathlib import Path
from fastapi import HTTPException
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from sqlalchemy.orm import Session
from app.ai_stubs.draft_fields import missing_fields
from app.models import AuditLog, DraftStatus, DraftedDocument
from app.services.query_service import session_for
from app.services.storage_service import storage
from app.templates.drafts.registry import TEMPLATES

def template_for(doc_type: str) -> dict:
    template = TEMPLATES.get(doc_type)
    if not template: raise HTTPException(400, f'Unsupported document type: {doc_type}')
    return template

def response_for(draft: DraftedDocument) -> dict:
    missing = missing_fields(draft.doc_type, draft.collected_fields)
    return {'id': draft.id, 'doc_type': draft.doc_type, 'status': draft.draft_status.value, 'missing_fields': missing, 'download_url': f'/api/v1/draft/{draft.id}/download' if draft.final_file_ref else None}

def start(db: Session, doc_type: str, fields: dict, session_id: int | None, user_id: int | None) -> dict:
    template_for(doc_type); session = session_for(db, session_id, user_id)
    draft = DraftedDocument(session_id=session.id, doc_type=doc_type, collected_fields=fields, draft_status=DraftStatus.ready if not missing_fields(doc_type, fields) else DraftStatus.collecting)
    db.add(draft); db.commit(); db.refresh(draft); return response_for(draft)

def add_answer(db: Session, draft_id: int, fields: dict) -> dict:
    draft = db.get(DraftedDocument, draft_id)
    if not draft: raise HTTPException(404, 'Draft not found')
    draft.collected_fields = {**draft.collected_fields, **fields}
    draft.draft_status = DraftStatus.ready if not missing_fields(draft.doc_type, draft.collected_fields) else DraftStatus.collecting
    db.commit(); db.refresh(draft); return response_for(draft)

def generate(db: Session, draft_id: int, user_id: int | None) -> dict:
    draft = db.get(DraftedDocument, draft_id)
    if not draft: raise HTTPException(404, 'Draft not found')
    template = template_for(draft.doc_type); missing = [key for key in template['required_fields'] if not draft.collected_fields.get(key)]
    if missing: raise HTTPException(400, detail={'message': 'Required fields are missing', 'missing_fields': missing})
    values = {**draft.collected_fields, 'date': date.today().isoformat()}; text = template['body'].format(**values)
    pdf_path = Path(storage.save('drafts', f'{draft.id}_{draft.doc_type}.pdf', b''))
    pdf = canvas.Canvas(str(pdf_path), pagesize=A4); pdf.setTitle(template['title']); y = 800
    for line in [template['title'], '', *text.splitlines()]: pdf.drawString(50, y, line[:110]); y -= 20
    pdf.save()
    docx_path = pdf_path.with_suffix('.docx'); document = Document(); document.add_heading(template['title'], 0)
    for line in text.splitlines(): document.add_paragraph(line)
    document.save(docx_path)
    draft.final_file_ref = str(pdf_path); draft.draft_status = DraftStatus.generated; db.add(AuditLog(user_id=user_id, action='draft_generate', resource_type='draft', resource_id=str(draft.id))); db.commit(); db.refresh(draft); return response_for(draft)

def download(db: Session, draft_id: int) -> str:
    draft = db.get(DraftedDocument, draft_id)
    if not draft or not draft.final_file_ref: raise HTTPException(404, 'Generated draft not found')
    return draft.final_file_ref
